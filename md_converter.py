#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown文档转换工具
支持转换为HTML、Word和PDF格式
"""

import os
import re
import markdown
import argparse
from pathlib import Path
from jinja2 import Template
from datetime import datetime
import base64
import mimetypes

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import weasyprint
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class MarkdownConverter:
    def __init__(self, input_file, output_dir="output"):
        self.input_file = Path(input_file)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # 读取Markdown内容
        with open(self.input_file, 'r', encoding='utf-8') as f:
            self.md_content = f.read()
        
        # 配置Markdown扩展
        self.md = markdown.Markdown(
            extensions=[
                'markdown.extensions.extra',
                'markdown.extensions.codehilite',
                'markdown.extensions.toc',
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code'
            ],
            extension_configs={
                'markdown.extensions.codehilite': {
                    'css_class': 'highlight'
                },
                'markdown.extensions.toc': {
                    'toc_depth': 3
                }
            }
        )
    
    def process_images(self, content, for_html=True):
        """处理图片路径，将相对路径转换为绝对路径或嵌入base64"""
        def replace_image(match):
            alt_text = match.group(1)
            img_path = match.group(2)
            
            # 相对路径转换为绝对路径
            if not img_path.startswith(('http://', 'https://', 'data:')):
                abs_path = self.input_file.parent / img_path
                if abs_path.exists():
                    if for_html:
                        # 对于HTML，转换为base64编码
                        with open(abs_path, 'rb') as img_file:
                            img_data = img_file.read()
                            mime_type, _ = mimetypes.guess_type(str(abs_path))
                            if mime_type:
                                b64_data = base64.b64encode(img_data).decode()
                                img_path = f"data:{mime_type};base64,{b64_data}"
                    else:
                        # 对于其他格式，复制图片到输出目录
                        output_img_dir = self.output_dir / "images"
                        output_img_dir.mkdir(exist_ok=True)
                        output_img_path = output_img_dir / abs_path.name
                        if not output_img_path.exists():
                            import shutil
                            shutil.copy2(abs_path, output_img_path)
                        img_path = f"images/{abs_path.name}"
            
            return f"![{alt_text}]({img_path})"
        
        # 匹配Markdown图片语法
        pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        return re.sub(pattern, replace_image, content)
    
    def get_html_template(self):
        """获取HTML模板"""
        return """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body {
            font-family: "Microsoft YaHei", "PingFang SC", Arial, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
            color: #333;
        }
        
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 2em;
            margin-bottom: 1em;
            font-weight: 600;
        }
        
        h1 {
            font-size: 2.5em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            text-align: center;
            margin-bottom: 2em;
        }
        
        h2 {
            font-size: 2em;
            border-bottom: 2px solid #3498db;
            padding-bottom: 0.3em;
        }
        
        h3 {
            font-size: 1.5em;
            color: #2980b9;
        }
        
        h4 {
            font-size: 1.25em;
            color: #34495e;
        }
        
        p {
            margin-bottom: 1em;
            text-align: justify;
        }
        
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border-radius: 8px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
            background-color: #fff;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        
        th {
            background-color: #3498db;
            color: white;
            font-weight: 600;
        }
        
        tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        
        tr:hover {
            background-color: #e8f4f8;
        }
        
        code {
            background-color: #f8f9fa;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            color: #e74c3c;
        }
        
        pre {
            background-color: #2c3e50;
            color: #ecf0f1;
            padding: 20px;
            border-radius: 8px;
            overflow-x: auto;
            margin: 20px 0;
        }
        
        pre code {
            background-color: transparent;
            padding: 0;
            color: inherit;
        }
        
        blockquote {
            border-left: 4px solid #3498db;
            margin: 20px 0;
            padding: 10px 20px;
            background-color: #f8f9fa;
            font-style: italic;
        }
        
        ul, ol {
            margin-bottom: 1em;
            padding-left: 2em;
        }
        
        li {
            margin-bottom: 0.5em;
        }
        
        .toc {
            background-color: #f8f9fa;
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 20px;
            margin: 20px 0;
        }
        
        .toc ul {
            list-style-type: none;
            padding-left: 0;
        }
        
        .toc li {
            margin: 5px 0;
        }
        
        .toc a {
            text-decoration: none;
            color: #3498db;
        }
        
        .toc a:hover {
            text-decoration: underline;
        }
        
        .footer {
            margin-top: 3em;
            padding-top: 2em;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9em;
        }
        
        .highlight {
            background-color: #f8f9fa;
            border-radius: 8px;
            padding: 1em;
            margin: 1em 0;
        }
        
        @media print {
            body {
                padding: 0;
                background-color: white;
            }
            
            img {
                max-width: 100%;
                page-break-inside: avoid;
            }
            
            h1, h2, h3, h4, h5, h6 {
                page-break-after: avoid;
            }
            
            table, blockquote {
                page-break-inside: avoid;
            }
        }
        
        @media (max-width: 768px) {
            body {
                padding: 10px;
                font-size: 14px;
            }
            
            h1 {
                font-size: 2em;
            }
            
            h2 {
                font-size: 1.5em;
            }
            
            table {
                font-size: 12px;
            }
            
            th, td {
                padding: 8px;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        {{ content }}
        <div class="footer">
            <p>生成时间: {{ timestamp }}</p>
            <p>由 Markdown Converter 工具生成</p>
        </div>
    </div>
</body>
</html>"""
    
    def to_html(self):
        """转换为HTML格式"""
        # 处理图片
        processed_content = self.process_images(self.md_content, for_html=True)
        
        # 转换为HTML
        html_content = self.md.convert(processed_content)
        
        # 提取标题
        title = "研究文档"
        lines = self.md_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                title = line[2:].strip()
                break
        
        # 使用模板
        template = Template(self.get_html_template())
        final_html = template.render(
            title=title,
            content=html_content,
            timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )
        
        # 保存HTML文件
        output_file = self.output_dir / f"{self.input_file.stem}.html"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(final_html)
        
        print(f"HTML文件已生成: {output_file}")
        return output_file
    
    def to_word(self):
        """转换为Word格式"""
        if not DOCX_AVAILABLE:
            print("错误: 需要安装 python-docx 库才能导出Word文档")
            print("请运行: pip install python-docx")
            return None
        
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(12)
        
        # 处理图片
        processed_content = self.process_images(self.md_content, for_html=False)
        lines = processed_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                
                if level == 1:
                    heading = doc.add_heading(title_text, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_heading(title_text, level=min(level, 3))
            
            # 处理图片
            elif line.startswith('!['):
                # 简单的图片处理，实际应用中可能需要更复杂的解析
                match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if match:
                    img_path = match.group(2)
                    if not img_path.startswith(('http://', 'https://', 'data:')):
                        abs_path = self.input_file.parent / img_path
                        if abs_path.exists():
                            try:
                                doc.add_picture(str(abs_path), width=Inches(6))
                            except Exception as e:
                                doc.add_paragraph(f"[图片: {match.group(1)}]")
            
            # 处理表格（简化处理）
            elif '|' in line and not line.startswith('|'):
                # 这里需要更复杂的表格解析逻辑
                doc.add_paragraph(line)
            
            # 处理普通段落
            else:
                # 移除Markdown语法
                clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # 粗体
                clean_line = re.sub(r'\*([^*]+)\*', r'\1', clean_line)  # 斜体
                clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)  # 代码
                
                if clean_line:
                    doc.add_paragraph(clean_line)
        
        # 保存Word文件
        output_file = self.output_dir / f"{self.input_file.stem}.docx"
        doc.save(str(output_file))
        
        print(f"Word文件已生成: {output_file}")
        return output_file
    
    def to_pdf(self):
        """转换为PDF格式"""
        if not PDF_AVAILABLE:
            print("错误: 需要安装 weasyprint 库才能导出PDF文档")
            print("请运行: pip install weasyprint")
            return None
        
        # 首先生成HTML
        html_file = self.to_html()
        
        # 转换为PDF
        output_file = self.output_dir / f"{self.input_file.stem}.pdf"
        
        try:
            weasyprint.HTML(filename=str(html_file)).write_pdf(str(output_file))
            print(f"PDF文件已生成: {output_file}")
            return output_file
        except Exception as e:
            print(f"PDF生成失败: {e}")
            return None
    
    def convert_all(self):
        """转换为所有格式"""
        results = {}
        
        print("开始转换...")
        
        # 转换为HTML
        try:
            results['html'] = self.to_html()
        except Exception as e:
            print(f"HTML转换失败: {e}")
            results['html'] = None
        
        # 转换为Word
        try:
            results['word'] = self.to_word()
        except Exception as e:
            print(f"Word转换失败: {e}")
            results['word'] = None
        
        # 转换为PDF
        try:
            results['pdf'] = self.to_pdf()
        except Exception as e:
            print(f"PDF转换失败: {e}")
            results['pdf'] = None
        
        print("\n转换完成！")
        return results

def main():
    parser = argparse.ArgumentParser(description='Markdown文档转换工具')
    parser.add_argument('input', help='输入的Markdown文件')
    parser.add_argument('-o', '--output', default='output', help='输出目录 (默认: output)')
    parser.add_argument('-f', '--format', choices=['html', 'word', 'pdf', 'all'], 
                       default='all', help='输出格式 (默认: all)')
    
    args = parser.parse_args()
    
    if not Path(args.input).exists():
        print(f"错误: 文件 {args.input} 不存在")
        return
    
    converter = MarkdownConverter(args.input, args.output)
    
    if args.format == 'html':
        converter.to_html()
    elif args.format == 'word':
        converter.to_word()
    elif args.format == 'pdf':
        converter.to_pdf()
    else:
        converter.convert_all()

if __name__ == '__main__':
    main() 